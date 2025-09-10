import os
import io
import zipfile
import fitz  # PyMuPDF

from pdf_fill import save_pdf_temp, load_pdf_bytes, get_pdf_page_info, apply_text_overlays

from flask import Flask, render_template, request, send_file, redirect, jsonify
from PIL import Image, ImageDraw, ImageFont
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from docx import Document
from fpdf import FPDF
import pytesseract
from pdf2image import convert_from_bytes

# =========================
# App & Config
# =========================
app = Flask(__name__)

# Tesseract path: prefer ENV, fallback to your Windows path
TESS_PATH = os.getenv("TESSERACT_PATH", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
try:
    pytesseract.pytesseract.tesseract_cmd = TESS_PATH
except Exception:
    pass

# Optional upload size limit (50 MB)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50 MB


# =========================
# Helpers: Conversions
# =========================
def word_to_pdf_stream(docx_stream) -> io.BytesIO:
    """Very basic DOCX -> PDF (text only) using python-docx + FPDF."""
    doc = Document(docx_stream)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for para in doc.paragraphs:
        text = para.text or ""
        if not text.strip():
            pdf.ln(6)
        else:
            pdf.multi_cell(0, 8, text)

    out = io.BytesIO()
    pdf.output(out)
    out.seek(0)
    return out


def pdf_to_word_stream(pdf_bytes: bytes) -> io.BytesIO:
    """Extracts text from PDF and writes to a DOCX."""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    doc = Document()
    for page in reader.pages:
        txt = page.extract_text() or ""
        if txt.strip():
            for line in txt.splitlines():
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("")  # keep spacing

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def jpg_to_pdf_stream(image_streams) -> io.BytesIO:
    """
    Combine one or more images into a single PDF.
    image_streams: iterable of file-like image streams.
    """
    images = []
    for s in image_streams:
        img = Image.open(s).convert("RGB")
        images.append(img)

    if not images:
        raise ValueError("No images provided")

    out = io.BytesIO()
    if len(images) == 1:
        images[0].save(out, format="PDF")
    else:
        images[0].save(out, format="PDF", save_all=True, append_images=images[1:])
    out.seek(0)
    return out


def pdf_to_jpg_zip_stream(pdf_bytes: bytes, jpeg_quality: int = 85) -> io.BytesIO:
    """
    Convert all pages of a PDF to JPG and return an in-memory ZIP.
    """
    pil_pages = convert_from_bytes(pdf_bytes)
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for i, page in enumerate(pil_pages, start=1):
            page_rgb = page.convert("RGB")
            img_buf = io.BytesIO()
            page_rgb.save(img_buf, format="JPEG", quality=jpeg_quality, optimize=True)
            img_buf.seek(0)
            zf.writestr(f"page_{i}.jpg", img_buf.read())
    zip_buf.seek(0)
    return zip_buf


def jpg_to_png_stream(img_stream) -> io.BytesIO:
    img = Image.open(img_stream).convert("RGBA")
    out = io.BytesIO()
    img.save(out, format="PNG", optimize=True)
    out.seek(0)
    return out


def png_to_jpg_stream(img_stream, quality=90) -> io.BytesIO:
    img = Image.open(img_stream).convert("RGB")
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=quality, optimize=True)
    out.seek(0)
    return out


def merge_pdfs_stream(pdf_streams) -> io.BytesIO:
    merger = PdfMerger()
    for s in pdf_streams:
        s.seek(0)
        merger.append(s)
    out = io.BytesIO()
    merger.write(out)
    merger.close()
    out.seek(0)
    return out


def protect_pdf_stream(pdf_bytes: bytes, password: str) -> io.BytesIO:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    if password:
        writer.encrypt(password)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out


def remove_pdf_pages_stream(pdf_bytes: bytes, remove_pages_input: str) -> io.BytesIO:
    """
    remove_pages_input: e.g., "1,3,5-7"
    """
    to_remove = set()
    if remove_pages_input:
        parts = [p.strip() for p in remove_pages_input.split(",") if p.strip()]
        for p in parts:
            if "-" in p:
                a, b = p.split("-", 1)
                try:
                    a_i = int(a)
                    b_i = int(b)
                    for num in range(min(a_i, b_i), max(a_i, b_i) + 1):
                        to_remove.add(num)
                except Exception:
                    pass
            else:
                try:
                    to_remove.add(int(p))
                except Exception:
                    pass

    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    total = len(reader.pages)
    # Convert to zero-based
    remove_zero_based = {n - 1 for n in to_remove if 1 <= n <= total}

    for i in range(total):
        if i not in remove_zero_based:
            writer.add_page(reader.pages[i])

    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out


# =========================
# Helpers: OCR
# =========================
def ocr_from_image_stream(img_stream) -> str:
    image = Image.open(img_stream)
    return pytesseract.image_to_string(image)


def ocr_from_pdf_bytes(pdf_bytes: bytes) -> str:
    pages = convert_from_bytes(pdf_bytes)
    text = []
    for page in pages:
        text.append(pytesseract.image_to_string(page))
    return "\n".join(text)


# =========================
# Helpers: Compression
# =========================
def compress_image_stream(img_stream, quality: int) -> io.BytesIO:
    """
    Compress any image to JPEG with the given quality.
    """
    img = Image.open(img_stream).convert("RGB")
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=quality, optimize=True)
    out.seek(0)
    return out


def compress_pdf_bytes(pdf_bytes: bytes, dpi: int, jpeg_quality: int) -> io.BytesIO:
    """
    Rasterize each PDF page at `dpi` and re-embed as JPEG with `jpeg_quality`.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        pix = page.get_pixmap(dpi=dpi)
        mode = "RGB" if pix.alpha == 0 else "RGBA"
        pil = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
        img_buf = io.BytesIO()
        pil.convert("RGB").save(img_buf, format="JPEG", quality=jpeg_quality, optimize=True)
        img_buf.seek(0)
        rect = page.rect
        page.clean_contents()
        page.insert_image(rect, stream=img_buf.getvalue())
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    doc.close()
    return out


# =========================
# Helpers: Text Watermark
# =========================
def add_text_watermark_to_pdf(pdf_bytes: bytes, text: str,
                              font_size=48, opacity=0.3, rotation=45) -> io.BytesIO:
    """
    Cross-version safe: draw text with PIL -> rotate -> insert as image.
    Works on old/new PyMuPDF (no matrix arg needed).
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    for page in doc:
        rect = page.rect
        page_w, page_h = rect.width, rect.height

        # Pick a font (fallback if TTF missing)
        try:
            font = ImageFont.truetype("arial.ttf", size=font_size)
        except Exception:
            font = ImageFont.load_default()

        # Measure text
        tmp_img = Image.new("RGBA", (10, 10), (0, 0, 0, 0))
        tmp_draw = ImageDraw.Draw(tmp_img)
        l, t, r, b = tmp_draw.textbbox((0, 0), text, font=font)
        text_w, text_h = r - l, b - t

        # Draw text onto a transparent image with small padding
        pad = 20
        txt_img = Image.new("RGBA", (text_w + 2 * pad, text_h + 2 * pad), (0, 0, 0, 0))
        draw = ImageDraw.Draw(txt_img)
        alpha = max(0, min(255, int(opacity * 255)))
        draw.text((pad, pad), text, font=font, fill=(128, 128, 128, alpha))

        # Rotate to any angle
        rotated = txt_img.rotate(rotation, expand=True, resample=Image.BICUBIC)

        # Scale to ~60% of page width (keep aspect ratio)
        target_w = page_w * 0.6
        scale = target_w / rotated.width
        target_w_pts = target_w
        target_h_pts = rotated.height * scale

        # Center on page
        x0 = rect.x0 + (page_w - target_w_pts) / 2
        y0 = rect.y0 + (page_h - target_h_pts) / 2
        img_rect = fitz.Rect(x0, y0, x0 + target_w_pts, y0 + target_h_pts)

        # Encode rotated image as PNG (preserves transparency)
        buf = io.BytesIO()
        rotated.save(buf, format="PNG")
        buf.seek(0)

        # Insert as image
        page.insert_image(img_rect, stream=buf.getvalue(), keep_proportion=True, overlay=True)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    doc.close()
    return out


def add_text_watermark_to_image(img_stream, text: str) -> io.BytesIO:
    base = Image.open(img_stream).convert("RGBA")
    W, H = base.size
    layer = Image.new("RGBA", base.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(layer)

    # Choose font
    try:
        font = ImageFont.truetype("arial.ttf", size=max(24, int(min(W, H) * 0.05)))
    except Exception:
        font = ImageFont.load_default()

    bbox = draw.textbbox((0, 0), text, font=font)
    text_w, text_h = bbox[2] - bbox[0], bbox[3] - bbox[1]

    txt_img = Image.new("RGBA", (text_w + 20, text_h + 20), (0, 0, 0, 0))
    txt_draw = ImageDraw.Draw(txt_img)
    txt_draw.text((10, 10), text, font=font, fill=(255, 255, 255, 90))
    rotated = txt_img.rotate(30, expand=True)

    rx = (W - rotated.width) // 2
    ry = (H - rotated.height) // 2
    layer.alpha_composite(rotated, (rx, ry))

    out_img = Image.alpha_composite(base, layer)
    out = io.BytesIO()
    out_img.save(out, format="PNG")
    out.seek(0)
    return out


# =========================
# Routes: Home
# =========================
@app.route('/')
def index():
    return render_template('index.html')


# =========================
# Routes: Convert (existing features)
# =========================
@app.route('/convert', methods=['POST'])
def convert():
    conversion_type = request.form.get('conversion_type', '').strip()

    # Compression level mapping
    compression_level = request.form.get('compression_level', '').strip().lower()
    # defaults
    img_quality = 75
    pdf_dpi = 100
    pdf_jpeg_q = 75
    if compression_level:
        if compression_level in ("large", "high"):
            img_quality = 35
            pdf_dpi = 72
            pdf_jpeg_q = 35
        elif compression_level == "medium":
            img_quality = 55
            pdf_dpi = 100
            pdf_jpeg_q = 55
        elif compression_level in ("less", "low"):
            img_quality = 75
            pdf_dpi = 150
            pdf_jpeg_q = 75

    # Watermark (text only)
    watermark_text_value = (request.form.get('watermark_text_value') or "").strip()

    # Other fields
    password = request.form.get('password') or ""
    remove_pages_input = request.form.get('remove_pages_input') or ""

    files = request.files.getlist('file')
    if not files:
        return "No file uploaded", 400

    try:
        # by default, use the first file for 1-file operations
        first_file = files[0]
        fname = (first_file.filename or "").lower()

        # =======================
        # Multi-file operations
        # =======================
        if conversion_type == "merge_pdfs":
            # accept multiple PDFs
            streams = []
            for f in files:
                if not f.filename.lower().endswith(".pdf"):
                    return "All files must be PDFs for merging.", 400
                pdf_bytes = f.read()
                f.seek(0)
                streams.append(io.BytesIO(pdf_bytes))
            out = merge_pdfs_stream(streams)
            return send_file(out, as_attachment=True, download_name="merged.pdf")

        # =======================
        # Single-file operations
        # =======================
        # ---- OCR ----
        if conversion_type == "ocr":
            if fname.endswith((".png", ".jpg", ".jpeg")):
                text = ocr_from_image_stream(first_file.stream)
            elif fname.endswith(".pdf"):
                pdf_bytes = first_file.read()
                first_file.seek(0)
                text = ocr_from_pdf_bytes(pdf_bytes)
            else:
                return "Unsupported file format for OCR", 400
            out = io.BytesIO(text.encode("utf-8"))
            return send_file(out, as_attachment=True, download_name="ocr_output.txt")

        # ---- Compress ----
        if conversion_type == "compress":
            if fname.endswith((".png", ".jpg", ".jpeg")):
                out = compress_image_stream(first_file.stream, quality=img_quality)
                return send_file(out, as_attachment=True, download_name="compressed_image.jpg")
            elif fname.endswith(".pdf"):
                pdf_bytes = first_file.read()
                first_file.seek(0)
                out = compress_pdf_bytes(pdf_bytes, dpi=pdf_dpi, jpeg_quality=pdf_jpeg_q)
                return send_file(out, as_attachment=True, download_name="compressed_file.pdf")
            else:
                return "Unsupported file format for compression", 400

        # ---- Watermark (text only) ----
        if conversion_type == "watermark":
            if not watermark_text_value:
                return "Please provide watermark text.", 400

            if fname.endswith(".pdf"):
                pdf_bytes = first_file.read()
                first_file.seek(0)
                out = add_text_watermark_to_pdf(pdf_bytes, watermark_text_value)
                return send_file(out, as_attachment=True, download_name="watermarked_text.pdf")
            elif fname.endswith((".png", ".jpg", ".jpeg")):
                out = add_text_watermark_to_image(first_file.stream, watermark_text_value)
                return send_file(out, as_attachment=True, download_name="watermarked.png")
            else:
                return "Watermark option is only available for PDF or Image files", 400

        # ---- Protect PDF ----
        if conversion_type == "protect_pdf":
            if not fname.endswith(".pdf"):
                return "Please upload a PDF to protect.", 400
            pdf_bytes = first_file.read()
            first_file.seek(0)
            out = protect_pdf_stream(pdf_bytes, password=password)
            return send_file(out, as_attachment=True, download_name="protected.pdf")

        # ---- Remove PDF Pages ----
        if conversion_type == "remove_pages":
            if not fname.endswith(".pdf"):
                return "Please upload a PDF to modify.", 400
            pdf_bytes = first_file.read()
            first_file.seek(0)
            out = remove_pdf_pages_stream(pdf_bytes, remove_pages_input)
            return send_file(out, as_attachment=True, download_name="modified.pdf")

        # ---- Word -> PDF ----
        if conversion_type == "word_to_pdf":
            if not fname.endswith(".docx"):
                return "Please upload a .docx file.", 400
            out = word_to_pdf_stream(first_file.stream)
            return send_file(out, as_attachment=True, download_name="output.pdf")

        # ---- PDF -> Word ----
        if conversion_type == "pdf_to_word":
            if not fname.endswith(".pdf"):
                return "Please upload a PDF file.", 400
            pdf_bytes = first_file.read()
            first_file.seek(0)
            out = pdf_to_word_stream(pdf_bytes)
            return send_file(out, as_attachment=True, download_name="output.docx")

        # ---- JPG -> PDF ----
        if conversion_type == "jpg_to_pdf":
            if not fname.endswith((".jpg", ".jpeg", ".png")):
                return "Please upload an image (JPG/PNG).", 400
            image_streams = [f.stream for f in files if f.filename.lower().endswith((".jpg", ".jpeg", ".png"))]
            if not image_streams:
                return "No valid images found.", 400
            out = jpg_to_pdf_stream(image_streams)
            return send_file(out, as_attachment=True, download_name="output.pdf")

        # ---- PDF -> JPG (ZIP) ----
        if conversion_type == "pdf_to_jpg":
            if not fname.endswith(".pdf"):
                return "Please upload a PDF file.", 400
            pdf_bytes = first_file.read()
            first_file.seek(0)
            out = pdf_to_jpg_zip_stream(pdf_bytes, jpeg_quality=85)
            return send_file(out, as_attachment=True, download_name="pdf_pages.zip")

        # ---- JPG -> PNG ----
        if conversion_type == "jpg_to_png":
            if not fname.endswith((".jpg", ".jpeg")):
                return "Please upload a JPG/JPEG image.", 400
            out = jpg_to_png_stream(first_file.stream)
            return send_file(out, as_attachment=True, download_name="output.png")

        # ---- PNG -> JPG ----
        if conversion_type == "png_to_jpg":
            if not fname.endswith(".png"):
                return "Please upload a PNG image.", 400
            out = png_to_jpg_stream(first_file.stream, quality=90)
            return send_file(out, as_attachment=True, download_name="output.jpg")

        # Fallback
        return "Invalid conversion type", 400

    except Exception as e:
        return f"Error: {e}", 500


# =========================
# Routes: PDF Form Fill (Click-anywhere UI)
# =========================
@app.route("/formfill")
def formfill_home():
    return render_template("formfill_home.html")


@app.route("/formfill/upload", methods=["POST"])
def formfill_upload():
    file = request.files.get("pdf")
    if not file or not (file.filename or "").lower().endswith(".pdf"):
        return "No PDF uploaded", 400
    pdf_id = save_pdf_temp(file.read())
    return redirect(f"/formfill/editor/{pdf_id}")


@app.route("/formfill/editor/<pdf_id>")
def formfill_editor(pdf_id):
    # Provide page sizes for UI (PDF.js will render visually)
    info = get_pdf_page_info(pdf_id)
    return render_template("formfill_editor.html", pdf_id=pdf_id, pdf_info=info)


@app.route("/formfill/file/<pdf_id>")
def formfill_file(pdf_id):
    pdf_bytes = load_pdf_bytes(pdf_id)
    return send_file(
        io.BytesIO(pdf_bytes),
        download_name=f"{pdf_id}.pdf",
        mimetype="application/pdf"
    )


@app.route("/formfill/apply/<pdf_id>", methods=["POST"])
def formfill_apply(pdf_id):
    try:
        data = request.get_json(force=True, silent=False)
        overlays = data.get("overlays", [])
    except Exception:
        return jsonify({"error": "Invalid JSON"}), 400

    filled_pdf = apply_text_overlays(pdf_id, overlays)
    return send_file(
        io.BytesIO(filled_pdf),
        as_attachment=True,
        download_name=f"filled_{pdf_id}.pdf",
        mimetype="application/pdf"
    )


# =========================
# Run
# =========================
if __name__ == '__main__':
    app.run(debug=True)
